from io import BytesIO
from datetime import datetime

from docx import Document
from docx.shared import Pt
from docx.shared import RGBColor
from docx.shared import Inches

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION

from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def shade_cell(cell, color="D9EAD3"):

    tcPr = cell._tc.get_or_add_tcPr()

    shd = OxmlElement("w:shd")

    shd.set(qn("w:fill"), color)

    tcPr.append(shd)


def create_csv_document(content):

    if content is None:
        return None

    if not isinstance(content, dict):
        return None

    doc = Document()

    doc = Document()

    section = doc.sections[0]

    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)

    title = doc.add_heading(
        "SPECPILOT",
        level=0
    )

    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_heading(
        "CSV Validation Documentation",
        level=1
    )

    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    p = doc.add_paragraph()

    p.add_run("Project : ").bold = True

    p.add_run(
        content.get(
            "project_name",
            "N/A"
        )
    )

    p = doc.add_paragraph()

    p.add_run("Generated On : ").bold = True

    p.add_run(
        datetime.now().strftime("%d-%b-%Y")
    )

    doc.add_page_break()

    ##################################################
    # USER STORIES
    ##################################################

    heading = doc.add_heading(
        "1. USER STORIES",
        level=1
    )

    table = doc.add_table(
        rows=1,
        cols=4
    )

    table.style = "Table Grid"

    hdr = table.rows[0].cells

    hdr[0].text = "ID"
    hdr[1].text = "Title"
    hdr[2].text = "Description"
    hdr[3].text = "Priority"

    for cell in hdr:

        shade_cell(cell)

        for run in cell.paragraphs[0].runs:
            run.bold = True

    for story in content.get("user_stories", []):

        row = table.add_row().cells

        row[0].text = story.get("id", "")

        row[1].text = story.get("title", "")

        row[2].text = story.get("description", "")

        row[3].text = story.get("priority", "")

    ##################################################
    # RISK ASSESSMENT
    ##################################################

    doc.add_page_break()

    doc.add_heading(
        "2. RISK ASSESSMENT",
        level=1
    )

    risk_table = doc.add_table(
        rows=1,
        cols=5
    )

    risk_table.style = "Table Grid"

    hdr = risk_table.rows[0].cells

    hdr[0].text = "Risk ID"
    hdr[1].text = "Title"
    hdr[2].text = "Business Impact"
    hdr[3].text = "Probability"
    hdr[4].text = "Risk Rating"

    for cell in hdr:

        shade_cell(cell)

        for run in cell.paragraphs[0].runs:
            run.bold = True

    for risk in content.get("risk_assessment", []):

        row = risk_table.add_row().cells

        row[0].text = risk.get("id", "")

        row[1].text = risk.get("title", "")

        row[2].text = risk.get("business_impact", "")

        row[3].text = risk.get("probability", "")

        row[4].text = risk.get("risk_rating", "")

    doc.add_paragraph()

    doc.add_heading(
        "Risk Mitigation Details",
        level=2
    )

    for risk in content.get("risk_assessment", []):

        p = doc.add_paragraph()

        p.add_run(risk.get("id", "") + " - ").bold = True

        p.add_run(risk.get("title", ""))

        doc.add_paragraph(
            "Description: " + risk.get("description", "")
        )

        doc.add_paragraph(
            "Mitigation: " + risk.get("mitigation", "")
        )

        doc.add_paragraph()

    ##################################################
    # TEST CASES
    ##################################################

    doc.add_page_break()

    doc.add_heading(
        "3. VALIDATION TEST CASES",
        level=1
    )

    for tc in content.get("test_cases", []):

      doc.add_heading(
        tc.get("id", ""),
        level=2
    )

      table = doc.add_table(
        rows=0,
        cols=2
    )

      table.style = "Table Grid"

      def add_row(key, value):

        row = table.add_row().cells

        row[0].text = str(key)

        row[1].text = str(value)

        shade_cell(row[0], "D9EAD3")

        for run in row[0].paragraphs[0].runs:
            run.bold = True

      add_row("Title", tc.get("title", ""))

      add_row("Objective", tc.get("objective", ""))

      add_row("Preconditions", tc.get("preconditions", ""))

    # -----------------------
    # Test Steps
    # -----------------------

    steps_data = tc.get("test_steps", "")

    if isinstance(steps_data, list):
        steps = "\n".join(
            f"{i+1}. {step}"
            for i, step in enumerate(steps_data)
        )
    else:
        steps = str(steps_data)

    add_row("Test Steps", steps)

    # -----------------------
    # Expected Results
    # -----------------------

    expected_data = tc.get("expected_results", "")

    if isinstance(expected_data, list):
        expected = "\n".join(
            f"{i+1}. {step}"
            for i, step in enumerate(expected_data)
        )
    else:
        expected = str(expected_data)

    add_row("Expected Results", expected)

    doc.add_paragraph()

    ##################################################
    # END
    ##################################################

    output = BytesIO()

    doc.save(output)

    output.seek(0)

    return output